
import os
from docx import Document

def extract_text_from_docx(docx_path, output_path):
    print(f"Extracting text from: {docx_path}")
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        
        print(f"Successfully wrote content to: {output_path}")
    except Exception as e:
        print(f"Error extracting text: {e}")

if __name__ == "__main__":
    source_file = "/Users/hardentie/Downloads/vscode/03课：地点代词、选择疑问词数字100以上.doc"
    # Note: python-docx expects .docx format. The 'file' command said it is Word 2007+, 
    # so it is likely a docx with a doc extension.
    # We might need to rename it temporarily or just try opening it.
    # python-docx checks magic numbers usually, but extension matching might be a thing.
    
    output_file = "/Users/hardentie/Downloads/vscode/learning/japanese/lesson3_doc_content.txt"
    extract_text_from_docx(source_file, output_file)
